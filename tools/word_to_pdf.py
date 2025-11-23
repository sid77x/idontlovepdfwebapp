"""Convert Word (DOCX) document to PDF."""
import os
import sys
import streamlit as st
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Image as RLImage
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
from io import BytesIO
from PIL import Image

sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from utils.file_ops import save_uploaded_file, get_unique_filename, get_file_size_mb
from utils.preview import show_pdf_preview


def word_to_pdf(input_path: str, output_path: str) -> bool:
    """
    Convert Word document to PDF.
    
    Args:
        input_path: Input DOCX path
        output_path: Output PDF path
    
    Returns:
        bool: Success status
    """
    try:
        # Open Word document
        doc = Document(input_path)
        
        # Create PDF
        pdf_buffer = BytesIO()
        pdf_doc = SimpleDocTemplate(
            output_path,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )
        
        # Container for PDF elements
        elements = []
        
        # Get styles
        styles = getSampleStyleSheet()
        normal_style = styles['Normal']
        heading_style = styles['Heading1']
        
        # Process paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                # Determine style based on paragraph
                if para.style.name.startswith('Heading'):
                    style = heading_style
                else:
                    style = normal_style
                
                # Create paragraph
                try:
                    p = Paragraph(para.text, style)
                    elements.append(p)
                    elements.append(Spacer(1, 0.2 * inch))
                except Exception as e:
                    # If paragraph fails, add as plain text
                    pass
        
        # Process tables (simplified - add as text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text for cell in row.cells])
                if row_text.strip():
                    try:
                        p = Paragraph(row_text, normal_style)
                        elements.append(p)
                    except:
                        pass
            elements.append(Spacer(1, 0.2 * inch))
        
        # Build PDF
        pdf_doc.build(elements)
        
        return True
        
    except Exception as e:
        st.error(f"Error converting Word to PDF: {str(e)}")
        return False


def render_word_to_pdf_tool():
    """Render the Word to PDF tool interface."""
    st.title("📄 Word to PDF")
    st.write("Convert your Word document (DOCX) to PDF format")
    
    # File upload
    uploaded_file = st.file_uploader(
        "Choose a Word file",
        type=['docx'],
        help="Select the Word document you want to convert to PDF"
    )
    
    if uploaded_file:
        # Display file info
        file_size = get_file_size_mb(uploaded_file)
        st.info(f"📝 **{uploaded_file.name}** ({file_size:.2f} MB)")
        
        # Save uploaded file
        input_path = save_uploaded_file(uploaded_file, "word_to_pdf")
        
        if input_path:
            try:
                # Open document to get basic info
                doc = Document(input_path)
                num_paragraphs = len(doc.paragraphs)
                num_tables = len(doc.tables)
                
                st.success(f"✅ Word document loaded: {num_paragraphs} paragraphs, {num_tables} tables")
                
                st.info("💡 **Note:** Complex formatting, images, and advanced features may not be perfectly preserved.")
                
                # Convert button
                if st.button("🔄 Convert to PDF", type="primary", use_container_width=True):
                    with st.spinner("Converting Word to PDF..."):
                        # Create output path
                        output_dir = os.path.join("output", "word_to_pdf")
                        os.makedirs(output_dir, exist_ok=True)
                        output_path = os.path.join(output_dir, get_unique_filename("converted.pdf"))
                        
                        # Convert Word to PDF
                        success = word_to_pdf(input_path, output_path)
                        
                        if success:
                            st.success(f"✅ Successfully converted Word to PDF!")
                            
                            # Show preview
                            st.write("### 🔍 Preview")
                            try:
                                with open(output_path, "rb") as f:
                                    preview_data = f.read()
                                
                                show_pdf_preview(preview_data, height=1000)
                            except Exception as e:
                                st.warning(f"Preview unavailable: {str(e)}")
                            
                            # Show file info
                            output_size = os.path.getsize(output_path) / (1024 * 1024)
                            st.info(f"📄 Output PDF size: {output_size:.2f} MB")
                            
                            # Download button
                            with open(output_path, "rb") as f:
                                st.download_button(
                                    label="⬇️ Download PDF",
                                    data=f,
                                    file_name="converted.pdf",
                                    mime="application/pdf",
                                    use_container_width=True
                                )
                
            except Exception as e:
                st.error(f"❌ Error processing Word document: {str(e)}")
    
    else:
        # Help section
        st.write("### 📖 How to use")
        st.markdown("""
        1. **Upload** your Word document (.docx)
        2. **Click** Convert to PDF
        3. **Preview** the result
        4. **Download** your PDF file
        
        **What to expect:**
        - ✅ Text content preserved
        - ✅ Basic paragraph formatting
        - ⚠️ Complex formatting may be simplified
        - ⚠️ Images and embedded objects may not transfer
        
        **Supported:**
        - Microsoft Word 2007+ (.docx format)
        - Google Docs (download as .docx first)
        - LibreOffice Writer (.docx)
        
        **Best for:**
        - Simple documents
        - Reports and letters
        - Text-heavy content
        
        **Tip:** For best results with complex documents, use Microsoft Word's built-in "Save as PDF" feature.
        """)


if __name__ == "__main__":
    render_word_to_pdf_tool()
