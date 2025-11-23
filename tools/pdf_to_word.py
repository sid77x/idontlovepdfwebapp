"""Convert PDF to Word (DOCX) document."""
import os
import sys
import streamlit as st
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches, Pt
from io import BytesIO

sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from utils.file_ops import save_uploaded_file, get_unique_filename, get_file_size_mb


def pdf_to_word(input_path: str, output_path: str, include_images: bool = True) -> bool:
    """
    Convert PDF to Word document.
    
    Args:
        input_path: Input PDF path
        output_path: Output DOCX path
        include_images: Whether to extract and include images
    
    Returns:
        bool: Success status
    """
    try:
        # Open PDF
        pdf_document = fitz.open(input_path)
        
        # Create Word document
        doc = Document()
        
        # Process each page
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            
            # Add page number heading
            if page_num > 0:
                doc.add_page_break()
            
            # Extract text
            text = page.get_text()
            
            # Add text to document
            if text.strip():
                # Split into paragraphs
                paragraphs = text.split('\n')
                for para in paragraphs:
                    if para.strip():
                        p = doc.add_paragraph(para)
                        # Set font
                        for run in p.runs:
                            run.font.name = 'Calibri'
                            run.font.size = Pt(11)
            
            # Extract and add images if requested
            if include_images:
                image_list = page.get_images()
                
                for img_index, img in enumerate(image_list):
                    try:
                        xref = img[0]
                        base_image = pdf_document.extract_image(xref)
                        image_bytes = base_image["image"]
                        
                        # Add image to document
                        image_stream = BytesIO(image_bytes)
                        try:
                            doc.add_picture(image_stream, width=Inches(4))
                        except:
                            # If adding image fails, skip it
                            pass
                    except Exception as e:
                        # Skip problematic images
                        continue
        
        # Save document
        doc.save(output_path)
        pdf_document.close()
        
        return True
        
    except Exception as e:
        st.error(f"Error converting PDF to Word: {str(e)}")
        return False


def render_pdf_to_word_tool():
    """Render the PDF to Word tool interface."""
    st.title("📝 PDF to Word")
    st.write("Convert your PDF to an editable Word document (DOCX)")
    
    # File upload
    uploaded_file = st.file_uploader(
        "Choose a PDF file",
        type=['pdf'],
        help="Select the PDF file you want to convert to Word"
    )
    
    if uploaded_file:
        # Display file info
        file_size = get_file_size_mb(uploaded_file)
        st.info(f"📄 **{uploaded_file.name}** ({file_size:.2f} MB)")
        
        # Save uploaded file
        input_path = save_uploaded_file(uploaded_file, "pdf_to_word")
        
        if input_path:
            try:
                # Get page count
                pdf_doc = fitz.open(input_path)
                total_pages = len(pdf_doc)
                pdf_doc.close()
                
                st.success(f"✅ PDF loaded: {total_pages} pages")
                
                # Conversion settings
                st.write("### ⚙️ Conversion Settings")
                
                include_images = st.checkbox(
                    "Include images",
                    value=True,
                    help="Extract and include images from the PDF in the Word document"
                )
                
                st.info("💡 **Note:** Complex PDF formatting may not be perfectly preserved. The output will be plain text with images.")
                
                # Convert button
                if st.button("🔄 Convert to Word", type="primary", use_container_width=True):
                    with st.spinner("Converting PDF to Word... This may take a moment for large files."):
                        # Create output path
                        output_dir = os.path.join("output", "pdf_to_word")
                        os.makedirs(output_dir, exist_ok=True)
                        output_path = os.path.join(output_dir, get_unique_filename("converted.docx"))
                        
                        # Convert PDF to Word
                        success = pdf_to_word(input_path, output_path, include_images)
                        
                        if success:
                            st.success(f"✅ Successfully converted PDF to Word!")
                            
                            # Show file info
                            output_size = os.path.getsize(output_path) / (1024 * 1024)
                            st.info(f"📄 Output file size: {output_size:.2f} MB")
                            
                            # Download button
                            with open(output_path, "rb") as f:
                                st.download_button(
                                    label="⬇️ Download Word Document",
                                    data=f,
                                    file_name="converted.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    use_container_width=True
                                )
                            
                            st.success("💡 **Tip:** Open the downloaded file in Microsoft Word or Google Docs to edit.")
                
            except Exception as e:
                st.error(f"❌ Error processing PDF: {str(e)}")
    
    else:
        # Help section
        st.write("### 📖 How to use")
        st.markdown("""
        1. **Upload** your PDF file
        2. **Choose** whether to include images
        3. **Click** Convert to Word
        4. **Download** your editable Word document
        
        **What to expect:**
        - ✅ Text content extracted and formatted
        - ✅ Images extracted (if enabled)
        - ⚠️ Complex layouts may be simplified
        - ⚠️ Tables and advanced formatting may not be preserved
        
        **Best for:**
        - Text-heavy documents
        - Simple reports and letters
        - Documents you need to edit
        
        **Not recommended for:**
        - Forms with complex layouts
        - Documents with extensive tables
        - Highly formatted design documents
        """)


if __name__ == "__main__":
    render_pdf_to_word_tool()
